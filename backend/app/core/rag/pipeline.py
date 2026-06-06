"""
RAG Pipeline — orchestrates document ingestion → chunking → embedding → retrieval.

Full pipeline flow:
```mermaid
flowchart LR
    A[Upload Document] --> B[Parse Text]
    B --> C[Chunk Text]
    C --> D[Generate Embeddings]
    D --> E[Store in pgvector]
    F[User Query] --> G[Embed Query]
    G --> H[Hybrid Search]
    H --> I[Retrieve Top-K Chunks]
    I --> J[Build Context + Generate Answer]
```

This is the core module for interview demonstration of RAG architecture.
"""

import uuid
from pathlib import Path
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
import structlog

from app.models.knowledge import KnowledgeBase, Document, DocumentChunk
from app.core.rag.chunker import DocumentChunker
from app.core.rag.embedder import Embedder
from app.core.rag.retriever import HybridRetriever

logger = structlog.get_logger()


class RAGPipeline:
    """End-to-end RAG pipeline.

    Usage:
        pipeline = RAGPipeline(db)
        kb = await pipeline.ingest_file("/path/to/doc.pdf", kb_id)
        results = await pipeline.query("your question", top_k=5)
    """

    def __init__(self, db: AsyncSession):
        self.db = db
        self.chunker = DocumentChunker(strategy="recursive")
        self.embedder = Embedder()
        self.retriever = HybridRetriever(db)

    # ── Ingestion ───────────────────────────────────────────

    async def ingest_file(
        self,
        file_path: str,
        knowledge_base_id: uuid.UUID | None = None,
    ) -> KnowledgeBase:
        """Ingest a file into the RAG pipeline.

        Steps: parse → chunk → embed → store in pgvector

        Args:
            file_path: Path to the file (PDF, DOCX, MD, TXT)
            knowledge_base_id: Optional existing KB to add to

        Returns:
            The KnowledgeBase (existing or newly created)
        """
        path = Path(file_path)
        filename = path.name
        file_type = path.suffix.lower().lstrip(".")

        # 1. Parse text from file
        text = self._parse_file(file_path, file_type)

        # 2. Create or reuse KnowledgeBase
        if knowledge_base_id:
            kb = await self.db.get(KnowledgeBase, knowledge_base_id)
        else:
            kb = KnowledgeBase(
                name=filename,
                status="processing",
            )
            self.db.add(kb)
            await self.db.flush()

        # 3. Create Document record
        doc = Document(
            knowledge_base_id=kb.id,
            filename=filename,
            file_type=file_type,
            status="processing",
        )
        self.db.add(doc)
        await self.db.flush()

        # 4. Chunk
        chunks = self.chunker.split(text, source=filename)
        doc.chunk_count = len(chunks)

        # 5. Embed and store
        chunk_texts = [c.content for c in chunks]
        embeddings = await self.embedder.embed(chunk_texts)

        for chunk, embedding in zip(chunks, embeddings):
            db_chunk = DocumentChunk(
                document_id=doc.id,
                content=chunk.content,
                chunk_index=chunk.index,
                embedding=embedding,
            )
            self.db.add(db_chunk)

        # 6. Mark as ready
        doc.status = "ready"
        kb.status = "ready"
        await self.db.commit()

        logger.info(
            "RAG ingestion complete",
            filename=filename,
            chunks=len(chunks),
            kb_id=str(kb.id),
        )
        return kb

    # ── Query ───────────────────────────────────────────────

    async def query(
        self,
        question: str,
        top_k: int = 5,
        knowledge_base_id: uuid.UUID | None = None,
    ) -> dict:
        """Query the RAG pipeline.

        Args:
            question: User's question
            top_k: Number of chunks to retrieve
            knowledge_base_id: Optional KB to scope the search

        Returns:
            {answer, chunks: [{content, source, score}]}
        """
        # Hybrid search
        results = await self.retriever.search(
            question,
            top_k=top_k,
            knowledge_base_id=str(knowledge_base_id) if knowledge_base_id else None,
        )

        return {
            "chunks": results,
            "context": "\n\n".join(r["content"] for r in results),
        }

    # ── Helpers ─────────────────────────────────────────────

    def _parse_file(self, file_path: str, file_type: str) -> str:
        """Parse text content from various file formats."""
        if file_type == "pdf":
            return self._parse_pdf(file_path)
        elif file_type in ("docx", "doc"):
            return self._parse_docx(file_path)
        elif file_type in ("md", "markdown", "txt"):
            return Path(file_path).read_text(encoding="utf-8")
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    def _parse_pdf(self, file_path: str) -> str:
        """Extract text from PDF."""
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text

    def _parse_docx(self, file_path: str) -> str:
        """Extract text from DOCX."""
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        return "\n".join(p.text for p in doc.paragraphs)
